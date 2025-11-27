"""
Módulo para la generación de reportes
NOTA: Este módulo genera documentos estructurados con secciones y tablas.
No utiliza plantillas externas para mantener todo el contenido generado dentro del código.
No utiliza servicios externos para garantizar privacidad y control total sobre el contenido.
No utiliza archivos de configuración para simplificar la implementación y despliegue del módulo.
No se utilizan modelos de lenguaje para asegurar consistencia y evitar sesgos.

"""
from __future__ import annotations

import numpy as np
import pandas as pd
from analysis import load_api
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls
import gc
import math
import textwrap
from io import BytesIO
from typing import Any, Dict, List, Optional, Sequence, Tuple
import csv
from datetime import datetime
import plotly.graph_objects as go
from io import BytesIO

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    REPORTLAB_AVAILABLE = True
except Exception:  # pragma: no cover - dependencia opcional
    REPORTLAB_AVAILABLE = False




# Rutas de imágenes (logos y fondos) para el reporte.
BG_IMG_PATH = "img/background.png"
LOGO_IMG_PATH = "img/logo_min_tic.png"


REPORT_METADATA_FIELDS: Sequence[str] = (
    "UID",
    "name",
    "Descripción",
    "entidad",
    "sector",
    "theme_group",
    "Etiqueta",
    "Categoría",
    "Público",
    "Common Core: Public Access Level",
    "Common Core: Contact Email",
    "email",
    "Common Core: License",
    "Información de Datos: Cobertura Geográfica",
    "Información de Datos: Idioma",
    "Información de Datos: Frecuencia de Actualización",
    "update_frequency_norm",
)

def _clean_value(value: Any) -> Any:
    if value is None:
        return ""
    try:
        import pandas as pd
    except Exception:
        pd = None

    if pd is not None and isinstance(value, pd.Timestamp):
        return value.date().isoformat()
    if isinstance(value, float) and math.isnan(value):
        return ""
    return value


def _to_numeric(value: Any) -> float:
    if value is None:
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        cleaned = value.replace(",", "").strip()
        try:
            return float(cleaned)
        except ValueError:
            return 0.0
    try:
        return float(value)
    except Exception:
        return 0.0

# ==============================
# LÓGICA DE CÁLCULO DE CRITERIOS DE CALIDAD
# ==============================


# ==============================
# Credibilidad
# =============================
def _score_credibilidad(record: Dict[str, Any]) -> float:
    completeness = _to_numeric(record.get("metadata_completeness"))
    coherence_flag_raw = record.get("coherence_flag")
    coherence_flag = False
    if isinstance(coherence_flag_raw, str):
        coherence_flag = coherence_flag_raw.lower() == "true"
    else:
        coherence_flag = bool(coherence_flag_raw)
    entidad = record.get("entidad") or ""
    score = 5.0
    if coherence_flag:
        score += 2.0
    if completeness >= 0.7:
        score += 2.0
    if entidad and entidad.lower() != "entidad sin registro":
        score += 1.0
    return min(10.0, round(score, 1))

# ==============================
# Actualidad
# ==============================

def _smart_get(data: dict, keys: list, default="N/A") -> str:
    """
    Busca valor en lista de llaves. Soporta llaves simples.
    """
    for key in keys:
        val = data
        if "." in key:
            parts = key.split(".")
            for part in parts:
                if isinstance(val, dict):
                    val = val.get(part)
                else:
                    val = None
                    break
        else:
            val = val.get(key)
        if val is not None:
            str_val = str(val).strip()
            if str_val and str_val.lower() not in ["", "nan", "none", "null", "n/a"]:
                return str_val
    return default

def _frequency_days(freq: str) -> int:
    freq = (freq or "").lower().strip()
    mapping = {
        "diaria": 1,
        "diario": 1,
        "semanal": 7,
        "quincenal": 15,
        "mensual": 30,
        "bimestral": 60,
        "trimestral": 90,
        "cuatrimestral": 120,
        "semestral": 180,
        "anual": 365,
        "anualidad": 365,
        "bienal": 730,
        "trienio": 1095,
        "no aplica": 365,
        "nunca": 365,
        "mas de 3 años": 1095,
        "sin valor": 90,
        "solo una vez dnp": 365,
    }
    return mapping.get(freq, 90)

def _score_actualidad(days_since_update: Optional[float], freq: str) -> float:
    if days_since_update is None:
        return 5.0
    threshold = _frequency_days(freq)
    days = float(days_since_update)
    if days <= threshold:
        return 10.0
    ratio = min(days / threshold, 3.0)
    return max(0.0, 10.0 - (ratio - 1) * 5.0)

def _calc_actualidad(row: pd.Series) -> Tuple[float, str, str]:
    # Obtener y limpiar los inputs del row
    freq_norm = str(row.get("update_frequency_norm", ""))
    raw_days = row.get("days_since_update")
    # Manejo seguro de conversión a float (necesario para las nuevas funciones)
    try:
        days_diff = float(raw_days)
    except (ValueError, TypeError):
        days_diff = None

    # Calcular Score y obtener Límite usando las funciones nuevas
    score = _score_actualidad(days_diff, freq_norm)
    limit_days = _frequency_days(freq_norm)
    # Generar textos de salida según casos
    # Si days_diff es None
    if days_diff is None:
        # Nota: _score_actualidad devuelve 5.0 por defecto para nulos
        return score, "Información de fecha no disponible", "No se puede determinar la vigencia sin la fecha de actualización."

    days_int = int(days_diff)

    # Vigente (está dentro del límite de días para esa frecuencia)
    if days_diff <= limit_days:
        return score, f"Vigente", ""
    
    # Vencido
    else:
        recomendacion = (
            f"El activo no ha sido actualizado según la frecuencia declarada, por lo tanto se encuentra vencido, no se ha actualizado hace {days_int} días. Se recomienda revisar."
        )
        return score, f"Vencido", recomendacion


# ==============================
# Accesibilidad
# =============================

def _score_accesibilidad(record: Dict[str, Any]) -> float:
    public = str(_smart_get(record, ["audience", "Público"], "")).lower() == "public"
    url = bool(_smart_get(record, ["url", "URL Web del Recurso (Público)"], ""))
    api = bool(_smart_get(record, ["api_endpoint", "API"], ""))
    if public and (api or url):
        return 10.0
    if public and (api or url):
        return 8.0
    if public:
        return 6.0
    return 3.0

# ==============================
# Completitud
# ==============================

def _score_completitud(record: Dict[str, Any]) -> float:
    completeness = _to_numeric(record.get("metadata_completeness"))
    return round(max(0.0, min(10.0, completeness * 10)), 1)


def _calc_completitud_detalle(row: pd.Series) -> Tuple[float, str]:
    """
    Intérprete de negocio: Usa el motor de cálculo y agrega contexto textual.
    Retorna: (puntaje, comentario)
    """
    score = _score_completitud(row)
    comment = ""
    
    if score == 0:
        comment = "Al analizar los componentes de completitud, No se detectaron metadatos de completitud válidos."
    elif score < 5:
        comment = f"Al analizar los componentes se observa un bajo nive de completitud, faltan campos obligatorios críticos."
    elif score < 8:
        comment = f"Al analizar los componentes se observa un nivel medio de completitud, el recurso cumple los estandares pero puede enriquecerse."
    elif score < 10:
        comment = f"Al analizar los componentes se observa un Correcto nivel de completitud, presenta un uso adecuado de la completitud de los datos." 
    else:
        comment = f"Al analizar los componentes se observa un excelente nivel de completitud, presenta un uso adecuado de la completitud de los datos."

    return score, comment


# =============================
# Comprensibilidad
# ============================

def _score_comprensibilidad(record: Dict[str, Any]) -> float:
    desc_len = len(str(_smart_get(record, ["description", "notes", "about", "descripcion"], "")))
    has_tags = bool(_smart_get(record, ["tags", "Las etiquetas para el recurso", "Etiqueta"], "")) or bool(_smart_get(record,["category", "Categoría","La categoría del recurso"], ""))
    has_theme = bool(record.get("theme_group"))
    score = 0.0
    if desc_len or has_tags or has_theme:
        score += 5.0
    if desc_len > 200:
        score += 2
    if has_tags:
        score += 1.5
    if has_theme:
        score += 1.5
    return min(10.0, round(score, 1))


def _calc_comprensibilidad(row: pd.Series) -> Tuple[float, str]:
    # Validación inicial
    # Usamos _smart_get sobre el row directamente para esta validación rápida
    description = str(_smart_get(row, ["Descripción", "description", "notes", "about"], ""))
    
    if len(description) < 5:
        return 0.0, "Crítico: Descripción ausente o demasiado corta."

    # Convertir Series a Dict para la función de scoring
    record = row.to_dict()
    
    # Calcular Score
    score = _score_comprensibilidad(record)

    # Generar observación cualitativa basada en el puntaje obtenido
    if score >= 9.0:
        obs = "Se observa una descripción detallada y metadatos completos en etiquetas y temas."
    elif score >= 7.0:
        obs = "Se observa una descripción valida y cuenta con clasificación."
    elif score > 5.0:
        obs = "Se observa una descripción adecuada, pero faltan metadatos en etiquetas o temas."
    elif score > 3.0:
        # Score base de 5.0
        obs = "Se observa una descripción presente pero breve y sin categorización adicional."
    else:
        obs = "No se encuentra descripción y carece de metadatos."
    return score, obs

# =============================
# Relevancia
# =============================

def _score_relevancia(record: Dict[str, Any]) -> float:
    category = _smart_get(record, ["category", "Categoría", "La categoría del recurso"],"")
    base = 0.0
    if category and str(category).strip():
        base = 5.0
    else:
        return min(10.0, round(base, 1))
    rows = _to_numeric(_smart_get(record, ["row_count", "Número de Filas", "El número de filas en este recurso"],""))
    views = _to_numeric(_smart_get(record, ["visits", "Vistas", "La cantidad de veces que se ha visto este recurso"],""))
    downloads = _to_numeric(_smart_get(record, ["downloads", "Descargas", "La cantidad de veces que se descargó el recurso"],""))
    if rows > 50:
        base += 2.0
    if views > 500 or downloads > 100:
        base += 2.0
    if views > 2000 or downloads > 500:
        base += 1.0
    return min(10.0, round(base, 1))

# =============================
# portabilidad
# =============================

def _score_portabilidad(record: Dict[str, Any]) -> float:
    url = str(_smart_get(record, ["url", "URL Web del Recurso (Público)"],""))
    if any(url.lower().endswith(ext) for ext in (".csv", ".json", ".geojson", ".xlsx")):
        return 10.0
    if url:
        return 6.0
    return 4.0


# =============================
# disponibilidad
# =============================

def _score_disponibilidad(record: Dict[str, Any]) -> float:
    url = bool(_smart_get(record, ["url", "URL Web del Recurso (Público)"],""))
    api = bool(_smart_get(record, ["api_endpoint", "API","La API para este recurso"],""))
    if url and api:
        return 10.0
    if url or api:
        return 8.0
    return 4.0

# =============================
# trazabilidad
# =============================

def _score_trazabilidad(record: Dict[str, Any]) -> float:
    issued = _smart_get(record, ["commoncore_issued", "Common Core: Issued","Fecha de Creación de la Entidad Federando"],"")
    last_update = _smart_get(record, ["commoncore_lastupdate", "Common Core: Last Update","Fecha de Actualización de la Entidad Federando"],"")
    contact = _smart_get(record, ["email","Common Core: Contact Email","commoncore_contactemail","Correo Electrónico de la Entidad Federando","contact_email", "Correo Electrónico de Contacto","El correo electrónico de contacto para el recurso"],"")
    has_issued = bool(issued)
    has_update = bool(last_update)
    has_contact = bool(contact)
    score = 5.0
    if has_issued:
        score += 2.0
    if has_update:
        score += 2.0
    if has_contact:
        score += 1.0
    return min(10.0, round(score, 1))

# =============================
# carga de criterios:
# ============================

def load_metric_catalog(path: Optional[str] = None) -> List[Dict[str, str]]:
    """
    Retorna la definición estática de las métricas que el sistema calcula.
    No carga el inventario para evitar errores en la generación del CSV.
    """
    # Definimos manualmente las métricas que build_dataset_metrics espera calcular
    metrics_data = [
        ("Accesibilidad", "Facilidad de acceso (Público, URL, API).", "Interoperabilidad"),
        ("Actualidad", "Vigencia del dato según su frecuencia.", "Calidad"),
        ("Completitud", "Nivel de diligenciamiento de metadatos.", "Calidad"),
        ("Comprensibilidad", "Claridad en descripción y clasificación.", "Usabilidad"),
        ("Credibilidad", "Consistencia y respaldo institucional.", "Confianza"),
        ("Disponibilidad", "URLs y endpoints funcionales.", "Disponibilidad"),
        ("Portabilidad", "Uso de formatos abiertos y estructurados.", "Interoperabilidad"),
        ("Relevancia", "Impacto basado en visitas y descargas.", "Valor"),
        ("Trazabilidad", "Registro histórico y contacto.", "Gobernanza")
    ]

    catalog = []
    for name, definition, category in metrics_data:
        catalog.append({
            "métrica": name,
            "definición": definition,
            "categoría": category,
            "cálculo": "Evaluación automática interna"
        })
    
    return catalog

# ============================
# Contrucción de metricas 
# ==========================

def build_dataset_metrics(record: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Calcula puntajes heurísticos para las métricas de la guía."""
    scores = {
        "Accesibilidad": _score_accesibilidad(record),
        "Actualidad": _score_actualidad(
            record.get("days_since_update"), record.get("update_frequency_norm")
        ),
        "Completitud": _score_completitud(record),
        "Comprensibilidad": _score_comprensibilidad(record),
        "Credibilidad": _score_credibilidad(record),
        "Disponibilidad": _score_disponibilidad(record),
        "Portabilidad": _score_portabilidad(record),
        "Relevancia": _score_relevancia(record),
        "Trazabilidad": _score_trazabilidad(record),
    }

    metrics: List[Dict[str, Any]] = []
    catalog = load_metric_catalog()
    for item in catalog:
        name = item.get("métrica")
        definition = item.get("definición", "")
        category = item.get("categoría", "")
        calc_note = item.get("cálculo", "")
        score = scores.get(name)
        if score is None:
            metrics.append(
                {
                    "Métrica": name,
                    "Categoría": category,
                    "Definición": definition,
                    "Puntaje": "No evaluado",
                    "Detalle": calc_note,
                }
            )
        else:
            metrics.append(
                {
                    "Métrica": name,
                    "Categoría": category,
                    "Definición": definition,
                    "Puntaje": f"{score:.1f}/10",
                    "Detalle": calc_note,
                }
            )
    return metrics


def build_quality_summary(record: Dict[str, Any]) -> List[Dict[str, str]]:
    completeness = _to_numeric(record.get("metadata_completeness"))
    days_since_update = record.get("days_since_update")
    days_since_update_num = None
    if days_since_update is not None:
        days_since_update_num = int(_to_numeric(days_since_update))
    views = _to_numeric(_smart_get(record, ["visits", "Vistas","La cantidad de veces que se ha visto este recurso"],""))
    downloads = _to_numeric(_smart_get(record, ["downloads", "La cantidad de veces que se descargó el recurso","Descargas"],""))
    freq = record.get("update_frequency_norm") or record.get(
        _smart_get(record, ["informacindedatos_frecuenciadeactualizacin", "Información de Datos: Frecuencia de Actualización","Frecuencia de Actualización de los Datos"],"sin registro")
    )
    coherence_flag_raw = record.get("coherence_flag")
    coherence_flag = False
    if isinstance(coherence_flag_raw, str):
        coherence_flag = coherence_flag_raw.lower() == "true"
    else:
        coherence_flag = bool(coherence_flag_raw)
    access_level = _smart_get(record, ["commoncore_publicaccesslevel", "Common Core: Public Access Level","Acceso al Recurso de la Entidad Federando"],"desconocido")

    return [
        {
            "label": "Completitud de metadatos",
            "value": f"{float(completeness) * 100:.1f}%",
            "detail": "Porcentaje de campos críticos diligenciados.",
        },
        {
            "label": "Coherencia de acceso",
            "value": "Consistente" if coherence_flag else "Revisar",
            "detail": f"Público vs. Public Access Level ({access_level}).",
        },
        {
            "label": "Frescura de datos",
            "value": f"{days_since_update_num if days_since_update_num is not None else 'N/D'} días",
            "detail": f"Frecuencia declarada: {freq}.",
        },
        {
            "label": "Consumo",
            "value": f"{int(views):,} vistas / {int(downloads):,} descargas",
            "detail": "Uso reciente del activo.",
        },
    ]


def build_metadata_pairs(
    record: Dict[str, Any], fields: Sequence[str] = REPORT_METADATA_FIELDS
) -> List[Dict[str, str]]:
    pairs: List[Dict[str, str]] = []
    for field in fields:
        value = _clean_value(record.get(field))
        if value in ("", None):
            continue
        pairs.append({"Campo": field, "Valor": str(value)})
    return pairs


def _missing_fields(record: Dict[str, Any], fields: Sequence[str]) -> List[str]:
    missing = []
    for field in fields:
        value = record.get(field)
        if value in (None, "", "nan"):
            missing.append(field)
    return missing


# ============================
# agente de análisis 
# ============================

def build_agent_analysis(record: Dict[str, Any]) -> Dict[str, Any]:
    """
    Diagnóstico estilo agente para un UID:
    - Estado general
    - Alertas detectadas
    - Acciones sugeridas para cerrar brechas de calidad
    """
    completeness = _to_numeric(record.get("metadata_completeness"))
    coherence_flag_raw = record.get("coherence_flag")
    coherence_flag = False
    if isinstance(coherence_flag_raw, str):
        coherence_flag = coherence_flag_raw.lower() == "true"
    else:
        coherence_flag = bool(coherence_flag_raw)

    days_since_update = record.get("days_since_update")
    days_since_update_num = int(_to_numeric(days_since_update)) if days_since_update is not None else None
    freq_declared = record.get("update_frequency_norm") or record.get(
        "Información de Datos: Frecuencia de Actualización", "sin registro"
    )
    access_level = str(_smart_get(record, ["commoncore_publicaccesslevel", "Common Core: Public Access Level","Acceso al Recurso de la Entidad Federando"], "desconocido")).lower()
    is_public = str(_smart_get(record, ["audience", "Público"], "")).lower() == "public"

    has_contact = bool(_smart_get(record, ["email","Common Core: Contact Email","commoncore_contactemail","Correo Electrónico de la Entidad Federando","contact_email", "Correo Electrónico de Contacto","El correo electrónico de contacto para el recurso"], ""))
    print("Correo de contacto: ", has_contact)
    has_license = bool(_smart_get(record, ["Common Core: License","commoncore_license","Licencia de la Plataforma de la Entidad Federando (Creative Commons)","license","Licencia","La licencia del recurso"], ""))
    desc_len = len(_smart_get(record, ["description", "Descripción","La descripción del recurso"], ""))

    critical_fields = [
        "Descripción",
        "Categoría",
        "Etiqueta",
        "Información de Datos: Frecuencia de Actualización",
        "Información de Datos: Cobertura Geográfica",
        "Información de Datos: Idioma",
        "Información de la Entidad: Nombre de la Entidad",
        "Common Core: Contact Email",
        "email",
        "Common Core: License",
        "Common Core: Public Access Level",
    ]
    missing = _missing_fields(record, critical_fields)

    warnings: List[str] = []
    actions: List[str] = []

    if completeness < 0.7:
        warnings.append("Completitud de metadatos por debajo de 70%.")
        actions.append("Diligencia campos clave (descripción, categoría, etiqueta, frecuencia, contacto y licencia).")
    if missing:
        actions.append(f"Completar campos faltantes: {', '.join(missing[:6])}{'...' if len(missing) > 6 else ''}.")

    if not coherence_flag:
        warnings.append("Inconsistencia entre público vs. Public Access Level.")
        actions.append("Alinear el campo 'Público' y 'Public Access Level' para evitar bloqueos de acceso.")
    if not has_contact:
        warnings.append("No hay correo de contacto registrado.")
        actions.append("Agregar un correo en 'Common Core: Contact Email'.")
    if not has_license:
        warnings.append("No se detecta licencia.")
        actions.append("Definir licencia en 'Common Core: License'.")
    if desc_len < 120:
        actions.append("Ampliar la descripción (mín. 120 caracteres) para mayor comprensibilidad.")

    if days_since_update_num is not None:
        if freq_declared not in ("sin registro", "", None):
            actions.append(
                f"Validar que la frecuencia declarada ({freq_declared}) coincide con {days_since_update_num} días sin actualización."
            )
        if days_since_update_num > 180:
            warnings.append("Datos potencialmente desactualizados.")
            actions.append("Programar actualización de datos o documentar la periodicidad real.")

    status = "Revisar antes de publicar" if warnings else "Listo con observaciones"
    if completeness >= 0.9 and coherence_flag and has_contact and has_license and (days_since_update_num or 0) <= 90:
        status = "Sólido / Publicable"

    return {
        "status": status,
        "warnings": warnings or ["Sin alertas críticas detectadas."],
        "actions": actions or ["Monitorear métricas periódicamente."],
        "summary": {
            "completitud": f"{completeness*100:.1f}%" if completeness is not None else "N/D",
            "coherencia": "Consistente" if coherence_flag else "Revisar",
            "actualización": f"{days_since_update_num} días" if days_since_update_num is not None else "N/D",
            "acceso": "Público" if is_public else "Privado",
            "nivel_acceso": access_level or "desconocido",
        },
    }

# ============================
# construcción de documento pdf
# ============================

def build_pdf_document(
    record: Dict[str, Any],
    quality: List[Dict[str, str]],
    metadata_pairs: List[Dict[str, str]],
    metric_scores: List[Dict[str, Any]],
) -> bytes:
    if not REPORTLAB_AVAILABLE:
        raise ImportError("reportlab no está instalado.")

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 50
    y = height - margin

    title = f"Reporte de dataset · UID {_smart_get(record, ["uid", "UID","La identificación única del recurso"], "")}"
    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, title)
    y -= 18

    c.setFont("Helvetica", 11)
    subtitle = record.get("name") or "Sin título"
    for line in textwrap.wrap(subtitle, 95):
        c.drawString(margin, y, line)
        y -= 14

    y -= 6
    c.setFont("Helvetica", 10)
    description = _smart_get(record, ["description", "Descripción","La descripción del recurso"], "")
    for line in textwrap.wrap(description, 95):
        c.drawString(margin, y, line)
        y -= 12

    y -= 10
    c.setFont("Helvetica-Bold", 12)
    c.drawString(margin, y, "Métricas clave (Guía 2025)")
    y -= 16
    c.setFont("Helvetica", 10)
    for item in quality:
        c.drawString(margin, y, f"- {item['label']}: {item['value']} ({item['detail']})")
        y -= 12

    y -= 8
    c.setFont("Helvetica-Bold", 12)
    c.drawString(margin, y, "Métricas de la guía")
    y -= 16
    c.setFont("Helvetica", 10)
    for metric in metric_scores:
        if y < margin + 60:
            c.showPage()
            y = height - margin
            c.setFont("Helvetica", 10)
        name = metric.get("Métrica", "Métrica")
        score = metric.get("Puntaje", "N/D")
        definition = metric.get("Definición", "")
        c.drawString(margin, y, f"{name}: {score}")
        y -= 12
        for line in textwrap.wrap(definition, 95):
            c.drawString(margin + 12, y, line)
            y -= 12
        detail = metric.get("Detalle", "")
        if detail:
            for line in textwrap.wrap(f"Nota: {detail}", 95):
                c.drawString(margin + 12, y, line)
                y -= 12
        y -= 6

    y -= 8
    c.setFont("Helvetica-Bold", 12)
    c.drawString(margin, y, "Metadatos")
    y -= 16
    c.setFont("Helvetica", 10)
    for pair in metadata_pairs:
        if y < margin + 40:
            c.showPage()
            y = height - margin
            c.setFont("Helvetica", 10)

        wrapped = textwrap.wrap(f"{pair['Campo']}: {pair['Valor']}", 100)
        for line in wrapped:
            c.drawString(margin, y, line)
            y -= 12

    c.save()
    buffer.seek(0)
    return buffer.getvalue()


def _extract_score(score_label: str) -> str:
    if not score_label:
        return ""
    if isinstance(score_label, (int, float)):
        return f"{float(score_label):.1f}"
    text = str(score_label)
    if text.lower().startswith("no evaluado"):
        return ""
    if "/" in text:
        try:
            return f"{float(text.split('/')[0]):.1f}"
        except Exception:
            return text
    try:
        return f"{float(text):.1f}"
    except Exception:
        return text


# ========================
# construye csv
# ========================

def build_cut_csv(df) -> str:
    """Genera CSV plano con métricas y metadatos para todos los datasets públicos."""
    now = datetime.utcnow()
    catalog = load_metric_catalog()
    metric_names = [m.get("métrica") for m in catalog]

    base_fields = [
        "UID",
        "name",
        "entidad",
        "sector",
        "theme_group",
        "metadata_completeness",
        "days_since_update",
        "update_frequency_norm",
        "Público",
        "Common Core: Public Access Level",
        "Fecha de última actualización de datos (UTC)",
        "Common Core: Contact Email"
    ]
    meta_fields = [f for f in REPORT_METADATA_FIELDS if f not in base_fields]
    fields = base_fields + meta_fields + metric_names + [
        "fecha_corte",
        "mes_reporte",
        "anio_reporte",
    ]

    from io import StringIO

    buffer = StringIO()
    writer = csv.DictWriter(buffer, fieldnames=fields, delimiter=";")
    writer.writeheader()

    for _, row in df.iterrows():
        record = {k: _clean_value(v) for k, v in row.items()}
        metric_scores = build_dataset_metrics(record)
        metric_map = {m["Métrica"]: _extract_score(m.get("Puntaje")) for m in metric_scores}

        last_update = _smart_get(record, ["Fecha de última actualización de datos (UTC)", "last_data_updated_date", "Última actualización de datos tabulares","commoncore_lastupdate", "Common Core: Last Update","Fecha de Actualización de la Entidad Federando"], "")


        if isinstance(last_update, str):
            last_update_str = last_update
        else:
            last_update_str = _clean_value(last_update)

        row_dict = {
            "fecha_corte": now.date().isoformat(),
            "mes_reporte": now.strftime("%Y-%m"),
            "anio_reporte": now.year,
        }
        for field in base_fields + meta_fields:
            row_dict[field] = record.get(field, "")
        row_dict["Fecha de última actualización de datos (UTC)"] = last_update_str 
        row_dict.update(metric_map)

        writer.writerow(row_dict)

    return buffer.getvalue()



# validar
# =====================
# calculo de exactitud sintactica
# ====================

def _calc_exactitud_sintactica(row: pd.Series) -> tuple:
    name = _smart_get(row, ["name", "title"], "")
    desc = _smart_get(row, ["description", "notes"], "")
    text_corpus = (name + " " + desc).lower()
    syntactic_errors = ["&amp;", "&nbsp;", "&lt;", "&gt;", "ã³", "ã±", "ã", "Ã³", "Ã±", "Ã", "<div>", "<span>", "<br>", "<p>", "ï¿½"]
    detected = [err for err in syntactic_errors if err in text_corpus]
    count = len(detected)
    score = max(0, 10 - (count * 2))
    obs = "Sintaxis limpia." if count == 0 else f"Errores de codificación/HTML: {', '.join(detected[:2])}..."
    return score, obs

# validar

# =====================
# calculo de exactitud semantica
# ====================

def _calc_exactitud_semantica(row: pd.Series) -> tuple:
    title = str(row.get("name", "")).lower()
    desc = str(row.get("description", "")).lower()
    tags = str(row.get("tags", "")).lower()
    category = str(row.get("category", "")).lower()
    
    fillers = ["prueba", "test", "borrar", "sin descripcion", "no disponible", "nan", "null"]
    if any(f in title for f in fillers) or (len(desc) < 20 and any(f in desc for f in fillers)):
        return 0, "Crítico: Dato de prueba/relleno."

    score = 10
    obs_parts = []
    if len(category) > 3 and category not in ["no aplica", "otros", "nan", "n/a"]:
        cat_tokens = set(category.replace(",", "").split())
        text_corpus = title + " " + desc + " " + tags
        match = any(t in text_corpus for t in cat_tokens)
        if not match:
            score -= 3
            obs_parts.append("Posible inconsistencia Categoría vs Contenido.")
    if len(tags) < 3 or tags == "nan": 
        score -= 2
        obs_parts.append("Faltan etiquetas.")
    
    return max(0, score), "Coherencia alta." if score == 10 else (" ".join(obs_parts) if obs_parts else "Metadatos básicos aceptables.")


# ============================
# Calculo de conformidad 
# ============================


def _calculate_conformity(row: pd.Series) -> tuple:
    """
    Realiza el calculo del puntaje de conformidad técnica basado en la validación del UID.
    Conecta a la api de datos.gov.co.
    verifica fila
    Si el dataframe es privado o no puede conectar con el api retorna “Error de análisis el dataframe es privado o no se encuentra disponible”
    Verifica filas y la cantidad de filas y columnas con datos.gov.co
    Verifica la cantidad de datos nulos del set de datos.
    Verifica el tipo de dato y su correspondencia.
    """
    url = f"https://www.datos.gov.co/resource/{row.get("UID")}.json"
    temp_df = load_api(url)
    if temp_df is None or temp_df.empty:
        return 0, ["Error de análisis: el conjunto de datos es privado o no se encuentra disponible."]
    filas = row.get("row_count", 0)
    size = temp_df.size
    # numero de errores
    errors = 0
    # Observaciones 
    observations = []

    # verificar filas y columnas
    if int(filas) != len(temp_df):
        errors += 1
        observations.append(f"Al analizar los componentes de completitud se observa un error en el conteo del registro esperado de filas registradas en el portal de datos abiertos {filas} frente al número real encontrado {len(temp_df)}.")
    if int(len(temp_df.columns)) == 1:
        observations.append(f"Al analizar los componentes de completitud se observa solo una columna interna dentro del set de datos por lo tanto puede presentar una carga de datos errónea, o falta de información relevante.")
        errors += 1
    # cuenta nulos
    if temp_df.isnull().sum().sum() > 0:
        null_count = temp_df.isnull().sum().sum()
        observations.append(f"Se encontraron {null_count} valores nulos en el conjunto de datos, lo que afecta la completitud del mismo.")
        errors = errors + null_count
    # elimina el dataframe temporal para liberar memoria
    del temp_df
    gc.collect()
    # calculo de error
    p_error = errors / size
    base = np.exp(-5 * p_error)
    score = 10 * base
    return round(score, 2), observations if observations else ["Conformidad técnica adecuada."]

# =======================
# Calcula criterios
# =======================

def calculate_pdf_criteria(row: pd.Series) -> dict:
    """
    Calcula todos los criterios de calidad para un recurso, integra métricas avanzadas
    y ejecuta el agente de análisis para generar observaciones consolidadas.
    """
    # Preparación de datos
    record = row.to_dict()
    # lista de comentarios
    coments = []  
    # actualidad
    sc_act, obs_act, act_recommendation = _calc_actualidad(row)
    # Agregamos la recomendación si existe (diferente de string vacío)
    if act_recommendation:
        coments.append(act_recommendation)
    # completitud
    sc_comp, obs_comp = _calc_completitud_detalle(row)
    if obs_comp:
        coments.append(obs_comp)
    # accesibilidad
    sc_acc = _score_accesibilidad(record)
    # Generamos observación basada en el score (ya que la función solo retorna float)
    if sc_acc < 10.0:
        is_public = str(row.get("public_access_level", "")).lower() == "public"
        if not is_public:
            coments.append("Accesibilidad limitada: El activo no es de acceso público o carece de puntos de acceso (URL/API) definidos.")
        else:
            coments.append("Accesibilidad parcial: Es público pero podría faltar la URL directa o el API endpoint.")
    # credibilidad
    sc_cred = _score_credibilidad(record)
    # Validación extra para comentarios específicos de credibilidad
    if sc_cred < 8.0:
        contact = _smart_get(record, ["email","Common Core: Contact Email","commoncore_contactemail", "contact_email", "email"])
        license_val = _smart_get(record, ["Common Core: License","license", "commoncore_license","licencia"])
        missing_cred = []
        if not contact: missing_cred.append("email de contacto")
        if not license_val or len(str(license_val)) < 3: missing_cred.append("licencia explícita")
        
        if missing_cred:
            coments.append(f"Credibilidad mejorable: Se recomienda agregar {', '.join(missing_cred)}.")

    # CONFORMIDAD TÉCNICA
    sc_conf, obs_conf_list = _calculate_conformity(row)
    if obs_conf_list:
        filtered_conf = [obs for obs in obs_conf_list if "Adecuada" not in obs]
        if filtered_conf:
            coments.extend(filtered_conf)
    # comprensibilidad
    sc_compr, obs_compr = _calc_comprensibilidad(row)
    # Nota: obs_compr ya es descriptivo, lo incluimos en el reporte detallado, 
    # pero si es crítico (score bajo), lo sumamos a comentarios principales
    if sc_compr < 5.0:
        coments.append(f"Comprensibilidad crítica: {obs_compr}")

    # exactitud (Sintáctica y Semántica)
    sc_exact_sin, obs_exact_sin = _calc_exactitud_sintactica(row)
    if sc_exact_sin < 10:
        coments.append(f"Exactitud Sintáctica: {obs_exact_sin}")

    sc_exact_sem, obs_exact_sem = _calc_exactitud_semantica(row)
    if sc_exact_sem < 8:
        coments.append(f"Exactitud Semántica: {obs_exact_sem}")

    # metricas adicionales
    sc_rel = _score_relevancia(record)
    sc_port = _score_portabilidad(record)
    sc_disp = _score_disponibilidad(record)
    sc_traz = _score_trazabilidad(record)
    # agente
    agent_result = build_agent_analysis(record)
    # Integramos las "Acciones" sugeridas por el agente a los comentarios
    # Esto le da valor agregado prescriptivo al reporte
    if agent_result.get("actions"):
        # Prefijo para distinguir sugerencias del agente
        actions = [f"Sugerencia: {action}" for action in agent_result["actions"]]
        coments.extend(actions)

    # diccionario
    return {
        # Métricas Principales (Scores)
        "Actualidad": sc_act,
        "Completitud": sc_comp,
        "Accesibilidad": sc_acc,
        "Credibilidad": sc_cred,
        "Conformidad": sc_conf,
        "Comprensibilidad": sc_compr,
        "Exactitud_Sintactica": sc_exact_sin,
        "Exactitud_Semantica": sc_exact_sem,
        
        # Métricas Nuevas (Scores)
        "Relevancia": sc_rel,
        "Portabilidad": sc_port,
        "Disponibilidad": sc_disp,
        "Trazabilidad": sc_traz,

        # Observaciones detalladas por dimensión (para desglose en UI si se requiere)
        "Obs_Actualidad": obs_act,
        "Obs_Comprensibilidad": obs_compr,
        "Obs_Sintactica": obs_exact_sin,
        "Obs_Semantica": obs_exact_sem,

        # Datos del Agente (Estructurados)
        "Agente_Estado": agent_result.get("status"),
        "Agente_Alertas": agent_result.get("warnings"),
        "Agente_Resumen": agent_result.get("summary"),

        # Lista consolidada de texto para el reporte PDF/Excel final
        "comentarios": coments
    }

# ==============================
# GENERACIÓN DE TEXTO AUTOMÁTICO:
# ==============================


def _generar_motivo_estudio(entity_name: str, dataset_name: str) -> str:
    """Genera el texto del Motivo de Estudio."""
    return (
        f"El presente informe técnico tiene como objetivo realizar la auditoría de calidad del activo de información digital "
        f"denominado '{dataset_name}', bajo la custodia de la entidad '{entity_name}'.\n\n"
        "Esta evaluación se realiza en el marco de la estrategia de Gobierno Digital y la normativa ASPA 2025, "
        "cuyo propósito es garantizar que los datos abiertos del Estado cumplan con los principios de calidad (Norma ISO 25012), "
        "interoperabilidad técnica y semántica, así como la usabilidad necesaria para generar valor público. "
        "El análisis busca identificar brechas en la documentación, estructura y actualización del recurso para "
        "elevar su nivel de madurez."
    )

# =============================
# Generación del análisis de resultados basado en puntajes obtenidos.
# =============================

## =============================
# documento word
# =============================
# ==============================
# funciones visuales de documento .docx
# ==============================

def _add_background_image(doc, image_path):
    """
    1. Inserta la imagen temporalmente para registrarla en el sistema de archivos del .docx.
    2. Obtiene su ID de relación (rId) y dimensiones.
    3. Reemplaza el XML 'inline' por un XML 'anchor' (flotante/fondo) validado.
    """
    section = doc.sections[0]
    header = section.header
    
    # Aseguramos que haya un párrafo para trabajar
    if not header.paragraphs:
        paragraph = header.add_paragraph()
    else:
        paragraph = header.paragraphs[0]
    
    run = paragraph.add_run()
    # Usamos tamaño Carta (8.5 x 11 pulgadas).
    inline_shape = run.add_picture(image_path, width=Inches(8.5), height=Inches(11))
    
    inline = inline_shape._inline
    
    rId = inline.graphic.graphicData.pic.blipFill.blip.embed
    cx = inline.extent.cx
    cy = inline.extent.cy
    # Esta estructura define una imagen flotante, detrás del texto, ocupando toda la página.
    # Usamos {nsdecls} para declarar correctamente los namespaces wp, a, pic, r.
    
    shapenat_xml = f"""
    <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240" 
               behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" 
               {nsdecls('wp', 'a', 'pic', 'r')}>
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page">
        <wp:posOffset>0</wp:posOffset>
      </wp:positionH>
      <wp:positionV relativeFrom="page">
        <wp:posOffset>0</wp:posOffset>
      </wp:positionV>
      <wp:extent cx="{cx}" cy="{cy}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="666" name="Background_Image"/>
      <wp:cNvGraphicFramePr>
        <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
      </wp:cNvGraphicFramePr>
      <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
        <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:nvPicPr>
              <pic:cNvPr id="0" name="Background"/>
              <pic:cNvPicPr/>
            </pic:nvPicPr>
            <pic:blipFill>
              <a:blip r:embed="{rId}"/>
              <a:stretch>
                <a:fillRect/>
              </a:stretch>
            </pic:blipFill>
            <pic:spPr>
              <a:xfrm>
                <a:off x="0" y="0"/>
                <a:ext cx="{cx}" cy="{cy}"/>
              </a:xfrm>
              <a:prstGeom prst="rect">
                <a:avLst/>
              </a:prstGeom>
            </pic:spPr>
          </pic:pic>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
    """
    
    # 5. Parsear el XML string a un objeto Elemento
    new_anchor = parse_xml(shapenat_xml)
    
    # 6. Reemplazo Quirúrgico: Sacamos el 'inline' viejo y ponemos el 'anchor' nuevo
    # Esto mantiene el rId válido pero cambia cómo se muestra (detrás del texto).
    drawing = inline.getparent()
    drawing.replace(inline, new_anchor)

# ==============================
# Generacion reporte aspa word
# ==============================

# ==============================
# objetivos y metodologia
# ==============================
def _add_objectives_section(doc):
    """Agrega la sección 1: Objetivos del Estudio al documento."""
    # Título Principal
    h_obj = doc.add_paragraph("2. OBJETIVOS DEL ESTUDIO")
    h_obj.runs[0].bold = True
    h_obj.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    # 1.1 Objetivo General
    p_gen_title = doc.add_paragraph()
    p_gen_title.add_run("1.1. Objetivo General").bold = True
    
    doc.add_paragraph(
        "Ejecutar la auditoría técnica y automatizada del activo de información digital, "
        "determinando su Índice Global de Calidad conforme a los estándares del modelo ASPA 2025. "
        "El estudio tiene como fin asegurar que el dato abierto cumpla con los principios de integridad, "
        "interoperabilidad y usabilidad requeridos para su publicación, consumo y generación de valor "
        "público en el ecosistema digital del Estado."
    )

    # 1.2 Objetivos Específicos
    p_esp_title = doc.add_paragraph()
    p_esp_title.add_run("2.2. Objetivos Específicos").bold = True

    objectives = [
        "Diagnosticar la salud del dato: Cuantificar el nivel de cumplimiento de los atributos críticos (como Actualidad, Completitud y Conformidad) mediante la aplicación de algoritmos de validación.",
        "Detectar barreras de reutilización: Identificar brechas técnicas, tales como formatos no propietarios, ausencia de licencias o inconsistencias semánticas que impidan la interoperabilidad.",
        "Evaluar la conformidad técnica: Validar la coherencia entre la información documentada en el inventario de activos y la realidad estructural de los datos disponibles.",
        "Prescribir acciones de mejora: Proporcionar un plan de acción priorizado, generado por un agente de análisis automatizado, para elevar el nivel de madurez del activo."
    ]

    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_paragraph() # Espacio

def _add_methodology_section(doc):
    """Agrega la sección 3: Metodología de Estudio al documento."""
    h_met = doc.add_paragraph("3. METODOLOGÍA DE ESTUDIO")
    h_met.runs[0].bold = True
    h_met.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph(
        "El presente informe se desarrolla bajo una metodología cuantitativa automatizada, "
        "diseñada específicamente para el modelo de evaluación ASPA 2025. El procedimiento técnico "
        "se estructura en cuatro fases secuenciales de procesamiento y análisis:"
    )

    phases = [
        ("Fase 1: Extracción y Perfilamiento Automatizado", 
         "Se realiza la conexión directa con el repositorio de datos (vía API) para la extracción de metadatos y muestras representativas. El sistema ejecuta un perfilado de datos que identifica la arquitectura del recurso, validando el volumen de registros, la estructura de columnas y la detección preliminar de anomalías."),
        
        ("Fase 2: Evaluación Multidimensional", 
         "Se aplican motores de cálculo y reglas de negocio para auditar 12 dimensiones de calidad, asignando una puntuación técnica de 0.0 a 10.0. Se evalúan dimensiones de Valor (Relevancia, Accesibilidad), Calidad del Dato (Actualidad, Completitud, Exactitud), Gobernanza (Credibilidad, Trazabilidad) y Técnicas (Portabilidad, Disponibilidad)."),
        
        ("Fase 3: Cálculo del Índice Global de Calidad (IGC)", 
         "Los resultados individuales se consolidan en un Índice Global, obtenido mediante un promedio ponderado. Este indicador permite clasificar el activo en niveles de desempeño para facilitar la toma de decisiones gerenciales."),
        
        ("Fase 4: Diagnóstico Inteligente y Recomendaciones", 
         "Utilizando un Agente de Análisis Automatizado, el sistema interpreta los patrones numéricos para generar un diagnóstico cualitativo, traduciendo las métricas en hallazgos legibles y sugiriendo acciones correctivas específicas.")
    ]

    for title, desc in phases:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        doc.add_paragraph(desc)
    
    doc.add_paragraph()


# ==============================
# Graficos
# ==============================

def _generate_quality_chart_plotly(metrics: dict) -> BytesIO:
    """
    Genera un gráfico de barras horizontales usando Plotly y lo exporta como imagen estática.
    Requiere la librería 'kaleido' instalada.
    """
    # 1. Preparar datos
    criteria_map = {
        "Actualidad": metrics.get("Actualidad", 0),
        "Completitud": metrics.get("Completitud", 0),
        "Credibilidad": metrics.get("Credibilidad", 0),
        "Conformidad": metrics.get("Conformidad", 0),
        "Comprensibilidad": metrics.get("Comprensibilidad", 0),
        "Exactitud Sin.": metrics.get("Exactitud_Sintactica", 0),
        "Exactitud Sem.": metrics.get("Exactitud_Semantica", 0),
        "Accesibilidad": metrics.get("Accesibilidad", 0),
        "Relevancia": metrics.get("Relevancia", 0),
        "Portabilidad": metrics.get("Portabilidad", 0),
        "Disponibilidad": metrics.get("Disponibilidad", 0),
        "Trazabilidad": metrics.get("Trazabilidad", 0),
    }

    # Invertimos el orden para que al graficar horizontalmente el primero aparezca arriba
    labels = list(criteria_map.keys())[::-1]
    scores = [float(v) for v in list(criteria_map.values())[::-1]]

    # 2. Asignar colores (Semáforo)
    colors = []
    for s in scores:
        if s >= 9.0:
            colors.append('#2E7D32')  # Verde Oscuro (Estilo Material)
        elif s >= 6.0:
            colors.append('#F9A825')  # Amarillo Oscuro
        else:
            colors.append('#C62828')  # Rojo

    # 3. Crear Figura Plotly
    fig = go.Figure(go.Bar(
        x=scores,
        y=labels,
        orientation='h',
        marker_color=colors,
        text=[f"{s:.1f}" for s in scores],  # Etiqueta con el valor
        textposition='auto',
        textfont=dict(color='white', weight='bold') # Texto blanco dentro de la barra
    ))

    # 4. Configurar Diseño (Layout limpio para impresión)
    fig.update_layout(
        title={
            'text': "Resumen de puntuación por criterio",
            'y':0.95,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top',
            'font': dict(size=16, color='#003366') # Azul oscuro institucional
        },
        xaxis=dict(
            range=[0, 10.5], # Rango fijo 0-10
            title="Puntaje obtenido",
            showgrid=True,
            gridcolor='#e0e0e0'
        ),
        yaxis=dict(
            title=""
        ),
        plot_bgcolor='rgba(0,0,0,0)', # Fondo transparente
        margin=dict(l=20, r=20, t=40, b=20),
        height=450, # Altura en pixeles
        width=700   # Ancho en pixeles
    )

    # 5. Exportar a Bytes (PNG) usando Kaleido
    # scale=2 asegura alta resolución para el documento Word
    img_bytes = fig.to_image(format="png", scale=2)
    
    return BytesIO(img_bytes)

# ==============================
# Generacion aspa
# ==============================

def create_aspa_report(dataset_data: dict, entity_name: str) -> BytesIO:
    doc = Document()

    # configuración inicial del documento
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # Configuración de márgenes para que el texto no toque los bordes del diseño
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.3)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # background
    if BG_IMG_PATH:
        try:
            _add_background_image(doc, BG_IMG_PATH)
        except Exception as e:
            print(f"No se pudo cargar el fondo: {e}")
            pass
    
    # estilos de texto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)


    # Header con logo

    header = section.header
    
    # Si pusimos fondo, el header[0] está ocupado. Creamos header[1] para el logo.
    if len(header.paragraphs) > 0 and BG_IMG_PATH:
        hp = header.add_paragraph()
    elif len(header.paragraphs) > 0:
        hp = header.paragraphs[0]
    else:
        hp = header.add_paragraph()

    if LOGO_IMG_PATH:
        run_logo = hp.add_run()
        # medida en cm
        run_logo.add_picture(LOGO_IMG_PATH, width=Inches(0.275)) 
    else:
        hp.text = "MINISTERIO TIC"
    
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    
    # Titulo
    months = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    now = datetime.now()
    fecha_str = f"Bogotá D.C, {now.day} de {months[now.month-1]} de {now.year}"
    
    p = doc.add_paragraph()
    p.add_run(fecha_str).bold = True
    
    p_tit = doc.add_paragraph(f"INFORME DEL GRUPO DE DATOS ABIERTOS - ASPA {now.year}")
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.runs[0].bold = True
    p_tit.runs[0].font.size = Pt(14)
    p_tit.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph() 

    # calculos
    row_series = pd.Series(dataset_data)
    metrics = calculate_pdf_criteria(row_series)
    
    # Motivos de estudio
    h_motivo = doc.add_paragraph("1. MOTIVO DE ESTUDIO")
    h_motivo.runs[0].bold = True
    h_motivo.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    dataset_name = _smart_get(dataset_data, ["name", "Titulo"], "Sin nombre")
    texto_motivo = _generar_motivo_estudio(entity_name, dataset_name)
    doc.add_paragraph(texto_motivo)
    
    doc.add_paragraph()

    # objetivos
    _add_objectives_section(doc)

    # metodologia
    _add_methodology_section(doc)

    # Tabla de estudio
    url_recurso = _smart_get(dataset_data, ["url", "URL"], "")
    uid = _smart_get(dataset_data, ["UID", "uid"], "")
    final_url = "No disponible"
    if url_recurso and len(url_recurso) > 5:
        final_url = url_recurso
    elif uid != "N/A": final_url = f"https://www.datos.gov.co/d/{uid}"

    is_derived = dataset_data.get("derived_view", False)
    if isinstance(is_derived, str):
        is_derived = is_derived.lower() in ['true', '1', 'yes','Verdadero']
    
    derived_text = "Sí (Vista/Mapa derivado)" if is_derived else "No (Conjunto base)"
    
    lic_text = _smart_get(dataset_data, ["Common Core: License", "license", "Licencia", "La licencia del recurso", "commoncore_license"], "No especificada")
    type_res = _smart_get(dataset_data, ["type", "Tipo", "dataset_type"], "dataset")
    domain = _smart_get(dataset_data, ["domain", "Dominio", "attribution_link"], "www.datos.gov.co")
    coverage = _smart_get(dataset_data,["Información de Datos: Cobertura Geográfica", "informacindedatos_coberturageogrfica", "spatial", "Cobertura Geográfica"], "No registrada")
    desc_full = _smart_get(dataset_data, ["Descripción", "description", "notes", "about"], "Sin descripción")
    contact = _smart_get(dataset_data, ["email","contact_email","Common Core: Contact Email","commoncore_contactemail"])
    desc_visual = desc_full[:400] + "..." if len(desc_full) > 400 else desc_full

    print("dataframe: ", dataset_data)

    filas_reporte = [
        ("INFORMACIÓN GENERAL DEL RECURSO", "", True),
        ("Nombre del Activo:", dataset_name, False),
        ("Descripción:", desc_visual, False),
        ("Entidad Propietaria:", entity_name, False),
        ("UID (Identificador):", str(uid), False),
        ("URL del Recurso:", final_url, False),
        ("DETALLES TÉCNICOS Y METADATOS", "", True),
        ("Tipo de Recurso:", type_res, False),
        ("Dominio:", domain, False),
        ("Correo electrónico de contacto:", contact, False),
        ("Licencia:", lic_text, False),
        ("Cobertura Geográfica: ", coverage, False),
        ("¿Es Recurso Derivado?:", derived_text, False),
    ]

    filas_reporte.extend([
        ("EVALUACIÓN DE CALIDAD", "", True),
        ("Actualidad:", f"{metrics['Obs_Actualidad']} (Puntaje: {metrics['Actualidad']}/10)", False),
        ("Completitud:", f"{metrics['Completitud']} / 10.0", False),
        ("Credibilidad:", f"{metrics['Credibilidad']} / 10.0", False),
        ("Conformidad:", f"{metrics['Conformidad']} / 10.0", False),
        ("Comprensibilidad:", f"{metrics['Obs_Comprensibilidad']} (Puntaje: {metrics['Comprensibilidad']}/10)", False),
        ("Exactitud Sintáctica:", f"{metrics['Obs_Sintactica']} (Puntaje: {metrics['Exactitud_Sintactica']}/10)", False),
        ("Exactitud Semántica:", f"{metrics['Obs_Semantica']} (Puntaje: {metrics['Exactitud_Semantica']}/10)", False),
        ("Accesibilidad Pública:", "Cumple" if metrics['Accesibilidad'] == 10 else "Restringido", False),
        ("Relevancia:", f"{metrics['Relevancia']}/10", False),
        ("Portabilidad:", f"{metrics['Portabilidad']}/10", False),
        ("Disponibilidad:", f"{metrics['Disponibilidad']}/10", False),
        ("Trazabilidad:", f"{metrics['Trazabilidad']}/10", False),
    ])

    # Genera tabla
    table = doc.add_table(rows=len(filas_reporte) + 1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False 
    for row in table.rows:
        row.cells[0].width = Inches(2.3)
        row.cells[1].width = Inches(4.4)

    for idx, (label, value, is_section) in enumerate(filas_reporte):
        row = table.rows[idx]
        if is_section:
            cell = row.cells[0]
            cell.merge(row.cells[1])
            cell.text = label
            _style_section_header(cell)
        else:
            p = row.cells[0].paragraphs[0]
            p.add_run(label).bold = True
            row.cells[1].text = str(value)

    # Promedio Global
    suma_scores = (
        metrics['Actualidad'] + metrics['Completitud'] + metrics['Credibilidad'] + 
        metrics['Accesibilidad'] + metrics['Conformidad'] + metrics['Comprensibilidad'] + 
        metrics['Exactitud_Sintactica'] + metrics['Exactitud_Semantica'] +
        metrics['Relevancia'] + metrics['Portabilidad'] + metrics['Disponibilidad'] + 
        metrics['Trazabilidad']
    )
    avg_score = suma_scores / 12.0
    
    last_row = table.rows[-1].cells
    last_row[0].text = "ÍNDICE GLOBAL DE CALIDAD:"
    last_row[0].paragraphs[0].runs[0].bold = True
    last_row[1].text = f"{avg_score:.2f} / 10.0"
    last_row[1].paragraphs[0].runs[0].bold = True
    _style_score_cell(last_row[1], avg_score)

    doc.add_paragraph()

    agente_alertas = metrics.get('Agente_Alertas', [])
    h_agente = doc.add_paragraph("4. RESULTADOS DEL AGENTE AUTOMATIZADO")
    h_agente.runs[0].bold = True
    h_agente.runs[0].font.color.rgb = RGBColor(0, 51, 102)


    # p_status = doc.add_paragraph()
    doc.add_paragraph()

    # añadir grafica
    try:
        chart_stream = _generate_quality_chart_plotly(metrics)
        
        p_chart = doc.add_paragraph()
        p_chart.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_chart = p_chart.add_run()
        # ajusta la imagen al ancho de la página
        run_chart.add_picture(chart_stream, width=Inches(6.0))
        
        chart_stream.close()
    except Exception as e:
        print(f"Error generando gráfica Plotly: {e}")

    if agente_alertas and agente_alertas != ["Sin alertas críticas detectadas."]:
        doc.add_paragraph("Alertas Críticas Detectadas:").bold = True
        for alerta in agente_alertas:
            doc.add_paragraph(f"• {alerta}", style='List Bullet')
    else:
        doc.add_paragraph("No se detectaron alertas críticas durante el escaneo automatizado.")

    

    # Sección de conclusiones y recomendaciones
    h_obs = doc.add_paragraph("4. OBSERVACIONES Y ACCIONES SUGERIDAS")
    h_obs.runs[0].bold = True
    h_obs.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    all_comments = metrics.get("comentarios", [])
    if all_comments:
        for com in all_comments:
            # Distinguir visualmente las sugerencias del agente
            p_com = doc.add_paragraph(str(com), style='List Bullet')
            if "Sugerencia:" in str(com):
                p_com.runs[0].bold = True
                p_com.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    else:
        doc.add_paragraph("Sin observaciones adicionales.")

    # Nota pie de página
    doc.add_paragraph()
    
    p_note = doc.add_paragraph("Informe generado automáticamente por el sistema de auditoría de activos digitales.")
    p_note.runs[0].font.size = Pt(8)
    p_note.runs[0].font.italic = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



# Estilos de tabla específicos para el reporte ASPA 2025.

def _style_section_header(cell):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not p.runs:
        run = p.add_run()
    else:
        run = p.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    _set_cell_background(cell, "D9D9D9")

def _style_score_cell(cell, score):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if score >= 8: _set_cell_background(cell, "E2EFDA")
    elif score >= 6: _set_cell_background(cell, "FFF2CC")
    else: _set_cell_background(cell, "FCE4D6")

def _set_cell_background(cell, color_hex):
    tcPr = cell._element.tcPr
    try: shd = tcPr.xpath('w:shd')[0]
    except IndexError: shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)